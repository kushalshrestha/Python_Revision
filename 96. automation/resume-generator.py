from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT


def add_horizontal_line(paragraph):
    p = paragraph._p  # get the lxml element of paragraph
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')        # single line
    bottom.set(qn('w:sz'), '12')              # line thickness (6 to 12 is typical)
    bottom.set(qn('w:space'), '1')            # space between text and border
    bottom.set(qn('w:color'), '000000')       # black color
    pbdr.append(bottom)
    pPr.append(pbdr)

def set_heading_style(doc):
    style = doc.styles['Heading 1']
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Set for Asian fonts as well (to avoid fallback)
    rPr = style.element.rPr
    rFonts = rPr.rFonts
    rFonts.set(qn('w:eastAsia'), "Arial")

def add_hyperlink(paragraph, url, display_text, font_name="Arial", font_size=10, bold=True):
    """Create a clickable hyperlink with formatting."""
    # Create relationship id
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create w:hyperlink tag
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create w:r element
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Font name
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)

    # Bold
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)

    # Font size
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_size * 2))  # in half-points
    rPr.append(sz)

    # Blue color
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    # Underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    # Text
    new_run.append(rPr)
    text = OxmlElement("w:t")
    text.text = display_text
    new_run.append(text)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_custom_margins(doc, top=1, bottom=1, left=1, right=1):
    """Set custom page margins in inches."""
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def set_default_font(doc, font_name="Arial", font_size=10):
    """Set default font style for the entire document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)

    # Set for Asian fonts as well
    rPr = style.element.rPr
    rFonts = rPr.rFonts
    rFonts.set(qn('w:eastAsia'), font_name)

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_centered_skills_paragraph(doc):
    text = ("Full Stack Development ∙ Cloud Computing ∙ SQL/ETL Development∙ Backend Development\n"
            "Microservices ∙ SOLID Principle ∙ Team Leadership ∙ Troubleshooting ∙ Debugging ∙ Dependency Injection")

    para = doc.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Fix for font name on some systems (sets eastAsia font)
    rFonts = run._element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'Arial')
    

def create_document():
    doc = Document()
    set_custom_margins(doc, top=0.5, bottom=0.5, left=0.5, right=0.5)
    set_default_font(doc, font_name="Arial", font_size=10)
    set_heading_style(doc)
    add_header(doc)
    add_summary(doc)
    add_skills(doc)
    add_experience(doc)
    add_education(doc)
    return doc


from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

def add_header(doc):
    # Main Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)

    run = name_para.add_run("KUSHAL SHRESTHA")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.bold = True


    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_after = Pt(4)

    subtitle_run = subtitle_para.add_run("Software Engineer | Data Engineer")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(12)
    subtitle_run.bold = True

    # Contact Info (centered)
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Phone (plain bold)
    phone_run = contact_para.add_run("+1(641) 233-9461 | ")
    phone_run.bold = True
    phone_run.font.name = "Arial"
    phone_run.font.size = Pt(10)

    # Email
    add_hyperlink(contact_para, "mailto:kushal8work@gmail.com", "kushal8work@gmail.com")

    # Separator
    contact_para.add_run(" | ").bold = True

    # LinkedIn (use your actual profile link if available)
    add_hyperlink(contact_para, "https://www.linkedin.com/in/kushal8work", "Linkedin")

    # Separator
    contact_para.add_run(" | ").bold = True

    # Website
    add_hyperlink(contact_para, "https://kushalshr.com", "https://kushalshr.com")



def add_summary(doc):
    heading = doc.add_heading("SUMMARY", level=1)
    add_horizontal_line(heading)
    summary = (
        "Software Engineer with around 9 years of professional experience in designing, developing, and leading highly scalable enterprise applications. "
        "Proficient in Java, Spring, Python, Flask, Web API, React, Microsoft SQL Server, Oracle and ETL development. Skilled in analyzing client requirements, planning software architecture, and writing robust code. "
        "Proven ability to integrate modules, APIs, and relational and non-relational databases, ensuring accuracy and performance. Proficient in automation, user levels, data layers, stored procedures, and interface design. "
        "Well-versed in unit testing and cloud computing with Amazon Web Services. Effective communication skills with cross-functional teams and stakeholders. "
        "Familiar with Agile Scrum and test-driven methodologies and industry experience in the US healthcare, government, and travel industry."
    )
    doc.add_paragraph(summary)

def add_skills(doc):
    heading = doc.add_heading("SKILLS", level=1)
    add_horizontal_line(heading)
    add_centered_skills_paragraph(doc)
    sections = {
        "Programming Languages": "Java, Python, JavaScript, TypeScript, T-SQL, PL/SQL",
        "Frameworks / Libraries / API": "Spring (MVC, Boot, JPA, Security, Cloud), Hibernate, Flask, SQLAlchemy, Pandas, NumPy, SpaCy, React, JDBC",
        "Web": "HTML, HTML5, CSS, jQuery, Bootstrap, Tailwind CSS",
        "Databases/Big Data": "Microsoft SQL Server, Oracle, MySQL, PostgreSQL, MongoDB, Redis, Cassandra, Snowflake, DynamoDB, RDS/Aurora",
        "Web Services / Cloud Infrastructure": "REST, SOAP, Microservices, AWS(S3, SNS, SQS, IAM, EC2, Lambda, Cloudwatch, Step Function, IAM, Redshift), Azure, GraphQL",
        "CI/CD Tools": "Github Actions, Jenkins, Docker, Kubernetes",
        "Big Data & Distributed Computing": "PySpark, Apache Kafka, Apache Airflow",
        "Tools": "Power BI, Jupyter Notebook, Papermill, Git, Swagger, JIRA, Trello, New Relic, SonarQube, JSON, XML, Apache Maven, JBoss/WildFly, Apache Tomcat, Prometheus, Grafana, Postman, Docker, JMeter, Vite, Splunk",
        "Other": "Waterfall, Agile, Scrum, Kanban, TDD, Data Warehousing, Performance Testing, Documentation, Design Patterns, Teamwork, Fast Paced, Code Optimization, Interpersonal skills, Vision, Problem-solving"
    }

    for heading, items in sections.items():
        para = doc.add_paragraph(style='List Bullet')
        run_heading = para.add_run(f"{heading}: ")
        run_heading.bold = True
        para.add_run(items)


def set_arial(paragraph, bold=False, italic=False, color=None):
    if not paragraph.runs:
        run = paragraph.add_run()
    else:
        run = paragraph.runs[0]

    font = run.font
    font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')  # Ensures compatibility for east Asian fonts

    font.size = Pt(10)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = color


def add_experience(doc):
    heading = doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
    add_horizontal_line(heading)

    experiences = [
        {
            "company": "CLACKAMAS COUNTY, Oregon City, OR, US",
            "duration": "March 2024 – Present",
            "title": "Senior Systems Project Analyst (Data Warehouse Developer)",
            "details": [
                "Led the architectural design and implemented a custom data warehouse to support business intelligence for HUD programs reporting.",
                "Developed data warehouse using Spring Boot, Python, Flask-RESTful, AWS, and React.js with Tailwind CSS.",
                "Implemented OAuth 2.0 SSO with MFA and a Python-based data profiling tool using Pandas.",
                "Built a PySpark deduplication pipeline reducing processing from 3+ hours to under 2 minutes.",
                "Designed data models in SQL Server for improved Power BI dashboard performance.",
                "Used Spring JPA Entities, DAOs, DTOs, and annotation-based DI and IoC."
            ],
            "tech": "Java, Spring Core, Spring AOP, Python, Flask, OAuth2, AWS, JUnit, Lombok, MS-Excel"
        },
        {
            "company": "BOOKMUNDI APS., Skovlunden, Denmark",
            "duration": "August 2021 – October 2022",
            "title": "Senior Software Engineer",
            "details": [
                "Led full SDLC and backend architecture for global travel portal using Spring Boot.",
                "Created analytics pipeline using AWS Glue, Snowflake, improving retention by 8%.",
                "Developed Python-based NLP destination tagger using SpaCy and AWS S3.",
                "Automated infra with CloudFormation, deployed via ECS Fargate and Kafka.",
                "Produced regular booking reports, enhancing stakeholder decision-making."
            ],
            "tech": "Java, Spring Boot, Hibernate, MySQL, Redis, Python, SpaCy, Kafka, AWS"
        },
        {
            "company": "COTIVITI INC., South Jordan, UT, US",
            "duration": "March 2019 – August 2021",
            "title": "Software Engineer",
            "details": [
                "Processed over 100 vendor datasets in the US healthcare system.",
                "Built Spring MVC REST API with secure authentication and exception handling.",
                "Mapped Oracle data models via Hibernate JPA ORM.",
                "Automated data movement with Oracle Data Integrator (ODI).",
                "Integrated S3, SQS/SNS, ALBs, and implemented unit/integration testing.",
                "Built Python automation tools for the AIPOPS team."
            ],
            "tech": "Spring Boot, Oracle, MicroStrategy, PL/SQL, ALB, ETL, ODI"
        },
        {
            "company": "SPIRALOGICS INC., Atlanta, GA, US",
            "duration": "July 2015 – March 2019",
            "title": "Software Engineer | Team Lead (2018-2019)",
            "details": [
                "Led MMSEA Section 111 automation for RREs, ensuring 98% compliance.",
                "Built backend in Spring Boot and frontend in ReactJS for Broadspire MA Reporter.",
                "Created/optimized SQL Server stored procedures and T-SQL code.",
                "Built AngularJS-based ACH transaction tool for US bank accounts.",
                "Wrote unit and integration tests in JUnit and Selenium."
            ],
            "tech": "Java, Spring Boot, AngularJS, SQL Server, SSIS, Selenium, JUnit, TestRail"
        }
    ]

    for exp in experiences:
        para = doc.add_paragraph()
        
        # Add a right-aligned tab stop at about 6 inches (adjust as needed)
        para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
        
        # Add the company name
        run_company = para.add_run(exp['company'])
        run_company.bold = True
        run_company.font.name = 'Arial'
        
        # Add tab character to jump to right-aligned tab stop
        para.add_run('\t')
        
        # Add the duration text
        run_duration = para.add_run(exp['duration'])
        run_duration.bold = True
        run_duration.font.name = 'Arial'
        set_arial(para, bold=True)

        # Title
        title_para = doc.add_paragraph(exp['title'])
        set_arial(title_para, italic=True)

        # Bullet Details
        for detail in exp['details']:
            detail_para = doc.add_paragraph(detail, style='List Bullet')
            set_arial(detail_para)

        # Technologies
        tech_para = doc.add_paragraph(f"Technologies Used: {exp['tech']}")
        set_arial(tech_para, italic=True)

def add_education(doc):
    heading = doc.add_heading("EDUCATION", level=1)
    add_horizontal_line(heading)

    educations = [
        "Maharishi International University, Fairfield, IA – Master of Science in Computer Science (Oct 2022 - Jan 2025)",
        "Advanced College of Engineering and Management, Lalitpur, Nepal – B.E. in Computer Engineering (2009 - 2013)"
    ]
    for edu in educations:
        doc.add_paragraph(edu, style='List Bullet')
        
        

def add_job_description_section(doc, job_description):
    heading = doc.add_heading("JOB DESCRIPTION", level=1)
    add_horizontal_line(heading)
    doc.add_paragraph(job_description)


if __name__ == "__main__":
    doc = create_document()
    doc.save("Kushal_Shrestha_Resume_Copy.docx")
